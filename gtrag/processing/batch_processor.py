"""
Batch processing module for multi-file indexing in gtrag system.

This module provides functionality for processing multiple documents in parallel,
supporting various file formats and providing progress tracking.
"""

import logging
import os
import re
from typing import List, Dict, Any, Optional, Union, Tuple
from dataclasses import dataclass, field
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import time
from datetime import datetime
import json

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    textract = None
    TEXTRACT_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class DocumentInfo:
    """Document information class"""
    file_path: str
    doc_id: str
    date: str
    content: str
    file_size: int
    file_format: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    processing_time: float = 0.0
    error: Optional[str] = None
    chunk_count: int = 0
    entity_count: int = 0
    relation_count: int = 0

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary format."""
        return {
            'file_path': self.file_path,
            'doc_id': self.doc_id,
            'date': self.date,
            'file_size': self.file_size,
            'file_format': self.file_format,
            'metadata': self.metadata,
            'processing_time': self.processing_time,
            'error': self.error,
            'chunk_count': self.chunk_count,
            'entity_count': self.entity_count,
            'relation_count': self.relation_count
        }


@dataclass
class BatchProcessingConfig:
    """Batch processing configuration"""
    max_workers: int = 4
    use_multiprocessing: bool = False
    batch_size: int = 10
    supported_formats: List[str] = field(default_factory=lambda: ['.txt', '.md', '.pdf', '.docx', '.json'])
    skip_on_error: bool = True
    save_progress: bool = True
    progress_file: str = "batch_progress.json"
    time_extraction_pattern: str = r'(20\d{2}Q[1-4]|Q[1-4]\s*20\d{2}|20\d{2}-\d{2}-\d{2}|20\d{2}-\d{2}|20\d{2}|(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+20\d{2})'


class FileParser:
    """File parser supporting multiple document formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_txt,
            '.md': self._parse_txt,
            '.json': self._parse_json,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
        }
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse file content and extract metadata"""
        path = Path(file_path)
        file_format = path.suffix.lower()
        
        if file_format not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_format}")
        
        try:
            parser = self.supported_formats[file_format]
            content, metadata = parser(file_path)
            
            return {
                'content': content,
                'metadata': metadata,
                'file_format': file_format,
                'file_size': path.stat().st_size
            }
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {e}")
            raise
    
    def _parse_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Parse text files"""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        content = ""
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if not content:
            raise ValueError(f"Could not decode file with any supported encoding: {file_path}")
        
        metadata = {
            'lines': content.count('\n') + 1, 
            'characters': len(content),
            'words': len(content.split())
        }
        return content, metadata
    
    def _parse_json(self, file_path: str) -> Tuple[str, Dict]:
        """Parse JSON files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Try to extract content from common JSON structures
        if isinstance(data, dict):
            content = data.get('content', data.get('text', data.get('body', str(data))))
        elif isinstance(data, list):
            content = '\n'.join(str(item) for item in data)
        else:
            content = str(data)
        
        metadata = {
            'json_keys': list(data.keys()) if isinstance(data, dict) else [],
            'json_type': type(data).__name__
        }
        return content, metadata
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Parse PDF files"""
        if not TEXTRACT_AVAILABLE:
            raise ImportError("textract library required for PDF parsing. Install with: pip install textract")
        
        try:
            content = textract.process(file_path).decode('utf-8')
            metadata = {
                'extracted_with': 'textract',
                'pages': content.count('\f') + 1  # Form feed character often indicates page breaks
            }
            return content, metadata
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Parse DOCX files"""
        if DOCX_AVAILABLE:
            try:
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                metadata = {
                    'paragraphs': len(doc.paragraphs),
                    'extracted_with': 'python-docx'
                }
                return content, metadata
            except Exception as e:
                logger.warning(f"Failed to parse DOCX with python-docx: {e}")
        
        # Fallback to textract
        if TEXTRACT_AVAILABLE:
            try:
                content = textract.process(file_path).decode('utf-8')
                metadata = {'extracted_with': 'textract'}
                return content, metadata
            except Exception as e:
                logger.error(f"Failed to parse DOCX with textract: {e}")
                raise
        
        raise ImportError("Either python-docx or textract required for DOCX parsing")


class TimeExtractor:
    """Extract time information from file names and content
    
    Supports multiple time formats:
    - Quarters: 2024Q1, Q1 2024
    - ISO dates: 2024-03-15, 2024-03
    - Years: 2024
    - Month names: March 2024, Mar 2024
    """
    
    def __init__(self, pattern: str = r'(20\d{2}Q[1-4]|Q[1-4]\s*20\d{2}|20\d{2}-\d{2}-\d{2}|20\d{2}-\d{2}|20\d{2}|(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+20\d{2})'):
        self.pattern = re.compile(pattern, re.IGNORECASE)
    
    def extract_time(self, file_path: str, content: str = "") -> Optional[str]:
        """Extract time information from file path or content
        
        Args:
            file_path: Path to the file
            content: File content to search
            
        Returns:
            Extracted time string in standardized format, or None if not found
        """
        # First try file path
        path_time = self._extract_from_text(str(file_path))
        if path_time:
            return path_time
        
        # Then try content (first 1000 characters)
        if content:
            content_time = self._extract_from_text(content[:1000])
            if content_time:
                return content_time
        
        return None
    
    def _extract_from_text(self, text: str) -> Optional[str]:
        """Extract time from text using regex"""
        matches = self.pattern.findall(text)
        if matches:
            time_str = matches[0].strip()
            # Normalize format
            if time_str.startswith('Q'):
                # Q1 2023 -> 2023Q1
                parts = time_str.split()
                if len(parts) == 2:
                    time_str = f"{parts[1]}{parts[0]}"
            return time_str.replace(' ', '')
        return None


class BatchProcessor:
    """Main batch processor for handling multiple documents"""
    
    def __init__(self, gtrag_system, config: Optional[BatchProcessingConfig] = None):
        """
        Initialize batch processor
        
        Args:
            gtrag_system: gtrag system instance
            config: Batch processing configuration
        """
        self.gtrag_system = gtrag_system
        self.config = config or BatchProcessingConfig()
        self.file_parser = FileParser()
        self.time_extractor = TimeExtractor(self.config.time_extraction_pattern)
        
        self.processed_documents: List[DocumentInfo] = []
        self.failed_documents: List[DocumentInfo] = []
    
    def process_directory(self, directory_path: str, 
                         file_pattern: str = "*",
                         recursive: bool = True) -> Dict[str, Any]:
        """
        Process all files in a directory
        
        Args:
            directory_path: Path to directory containing documents
            file_pattern: File pattern to match (e.g., "*.pdf")
            recursive: Whether to search subdirectories
            
        Returns:
            Processing results summary
        """
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        # Find files
        if recursive:
            files = list(directory_path.rglob(file_pattern))
        else:
            files = list(directory_path.glob(file_pattern))
        
        # Filter by supported formats
        supported_files = [
            f for f in files 
            if f.suffix.lower() in self.config.supported_formats
        ]
        
        logger.info(f"Found {len(supported_files)} supported files in {directory_path}")
        
        return self.process_files(supported_files)
    
    def process_files(self, file_paths: List[Union[str, Path]]) -> Dict[str, Any]:
        """
        Process a list of files
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Processing results summary
        """
        start_time = time.time()
        
        # Convert to Path objects
        file_paths = [Path(p) for p in file_paths]
        
        logger.info(f"Starting batch processing of {len(file_paths)} files")
        
        # Load progress if enabled
        if self.config.save_progress:
            self._load_progress()
        
        # Process files in batches
        total_processed = 0
        total_failed = 0
        
        for i in range(0, len(file_paths), self.config.batch_size):
            batch_files = file_paths[i:i + self.config.batch_size]
            
            if self.config.use_multiprocessing:
                batch_results = self._process_batch_multiprocess(batch_files)
            else:
                batch_results = self._process_batch_threaded(batch_files)
            
            # Update results
            for result in batch_results:
                if result.error:
                    self.failed_documents.append(result)
                    total_failed += 1
                else:
                    self.processed_documents.append(result)
                    total_processed += 1
            
            # Save progress
            if self.config.save_progress:
                self._save_progress()
            
            logger.info(f"Processed batch {i//self.config.batch_size + 1}: "
                       f"{len([r for r in batch_results if not r.error])} success, "
                       f"{len([r for r in batch_results if r.error])} failed")
        
        # Build temporal links after all documents are processed
        if total_processed > 0:
            logger.info("Building temporal links...")
            self.gtrag_system.build_temporal_links()
        
        total_time = time.time() - start_time
        
        results = {
            'total_files': len(file_paths),
            'processed_successfully': total_processed,
            'failed': total_failed,
            'processing_time': total_time,
            'avg_time_per_file': total_time / len(file_paths) if file_paths else 0,
            'processed_documents': [doc.to_dict() for doc in self.processed_documents],
            'failed_documents': [doc.to_dict() for doc in self.failed_documents]
        }
        
        logger.info(f"Batch processing completed: {total_processed} success, {total_failed} failed in {total_time:.2f}s")
        return results
    
    def _process_batch_threaded(self, file_paths: List[Path]) -> List[DocumentInfo]:
        """Process batch of files using threads"""
        results = []
        
        with ThreadPoolExecutor(max_workers=self.config.max_workers) as executor:
            future_to_file = {
                executor.submit(self._process_single_file, file_path): file_path
                for file_path in file_paths
            }
            
            for future in as_completed(future_to_file):
                result = future.result()
                results.append(result)
        
        return results
    
    def _process_batch_multiprocess(self, file_paths: List[Path]) -> List[DocumentInfo]:
        """Process batch of files using processes"""
        results = []
        
        with ProcessPoolExecutor(max_workers=self.config.max_workers) as executor:
            future_to_file = {
                executor.submit(self._process_single_file, file_path): file_path
                for file_path in file_paths
            }
            
            for future in as_completed(future_to_file):
                result = future.result()
                results.append(result)
        
        return results
    
    def _process_single_file(self, file_path: Path) -> DocumentInfo:
        """Process a single file"""
        start_time = time.time()
        doc_id = file_path.stem
        
        try:
            # Parse file
            parsed_data = self.file_parser.parse_file(str(file_path))
            content = parsed_data['content']
            
            # Extract time information
            date = self.time_extractor.extract_time(str(file_path), content)
            if not date:
                date = "UNKNOWN"
            
            # Prepare metadata
            metadata = {
                'date': date,
                'file_path': str(file_path),
                **parsed_data['metadata']
            }
            
            # Insert into gtrag system
            self.gtrag_system.insert(content, doc_id, metadata)
            
            processing_time = time.time() - start_time
            
            # Get processing statistics
            stats = self.gtrag_system.get_stats()
            
            return DocumentInfo(
                file_path=str(file_path),
                doc_id=doc_id,
                date=date,
                content=content[:200] + "..." if len(content) > 200 else content,
                file_size=parsed_data['file_size'],
                file_format=parsed_data['file_format'],
                metadata=metadata,
                processing_time=processing_time,
                chunk_count=stats.get('indexed_chunks', 0),
                # Note: entity/relation counts would need to be tracked per document
                entity_count=0,
                relation_count=0
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = str(e)
            
            logger.error(f"Failed to process {file_path}: {error_msg}")
            
            return DocumentInfo(
                file_path=str(file_path),
                doc_id=doc_id,
                date="UNKNOWN",
                content="",
                file_size=0,
                file_format=file_path.suffix.lower(),
                processing_time=processing_time,
                error=error_msg
            )
    
    def _save_progress(self):
        """Save processing progress to file"""
        progress_data = {
            'processed_documents': [doc.to_dict() for doc in self.processed_documents],
            'failed_documents': [doc.to_dict() for doc in self.failed_documents],
            'timestamp': datetime.now().isoformat()
        }
        
        with open(self.config.progress_file, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, indent=2, ensure_ascii=False)
    
    def _load_progress(self):
        """Load processing progress from file"""
        if not os.path.exists(self.config.progress_file):
            return
        
        try:
            with open(self.config.progress_file, 'r', encoding='utf-8') as f:
                progress_data = json.load(f)
            
            # Reconstruct document info objects
            self.processed_documents = [
                DocumentInfo(**doc_data) 
                for doc_data in progress_data.get('processed_documents', [])
            ]
            
            self.failed_documents = [
                DocumentInfo(**doc_data) 
                for doc_data in progress_data.get('failed_documents', [])
            ]
            
            logger.info(f"Loaded progress: {len(self.processed_documents)} processed, "
                       f"{len(self.failed_documents)} failed")
                       
        except Exception as e:
            logger.error(f"Failed to load progress: {e}")
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics"""
        if not self.processed_documents:
            return {}
        
        processing_times = [doc.processing_time for doc in self.processed_documents]
        file_sizes = [doc.file_size for doc in self.processed_documents]
        
        return {
            'total_processed': len(self.processed_documents),
            'total_failed': len(self.failed_documents),
            'avg_processing_time': sum(processing_times) / len(processing_times),
            'min_processing_time': min(processing_times),
            'max_processing_time': max(processing_times),
            'total_file_size': sum(file_sizes),
            'avg_file_size': sum(file_sizes) / len(file_sizes),
            'times_found': len(set(doc.date for doc in self.processed_documents)),
            'file_formats': list(set(doc.file_format for doc in self.processed_documents))
        }