"""
Batch processing module for multi-file indexing in TimeRAG system
"""
import logging
import os
import re
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import time
from datetime import datetime
import json


import textract


logger = logging.getLogger(__name__)


@dataclass
class DocumentInfo:
    """文檔信息類別"""
    file_path: str
    doc_id: str
    quarter: str
    content: str
    file_size: int
    file_format: str
    metadata: Dict[str, Any]
    processing_time: float = 0.0
    error: Optional[str] = None
    chunk_count: int = 0
    entity_count: int = 0
    relation_count: int = 0


@dataclass
class BatchProcessingConfig:
    """批量處理配置"""
    max_workers: int = 4
    use_multiprocessing: bool = False
    batch_size: int = 10
    supported_formats: List[str] = None
    skip_on_error: bool = True
    save_progress: bool = True
    progress_file: str = "batch_progress.json"


class FileParser:
    """文件解析器 - 支援多種文件格式"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._parse_txt,
            '.md': self._parse_txt,
            '.json': self._parse_json,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
        }
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """解析文件內容"""
        path = Path(file_path)
        file_format = path.suffix.lower()
        
        if file_format not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_format}")
        
        try:
            parser = self.supported_formats[file_format]
            content, metadata = parser(file_path)
            
            return {
                'content': content,
                'metadata': metadata,
                'file_format': file_format,
                'file_size': path.stat().st_size
            }
        except Exception as e:
            logger.error(f"解析文件失敗 {file_path}: {e}")
            raise
    
    def _parse_txt(self, file_path: str) -> tuple[str, Dict]:
        """解析文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        metadata = {'lines': content.count('\n') + 1, 'characters': len(content)}
        return content, metadata
    
    def _parse_json(self, file_path: str) -> tuple[str, Dict]:
        """解析JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        content = str(data.get('content', '') or data.get('text', '') or data)
        metadata = {'keys': list(data.keys()) if isinstance(data, dict) else []}
        return content, metadata
    
    def _parse_pdf(self, file_path: str) -> tuple[str, Dict]:
        """使用 textract 解析PDF文件。"""
        if not textract:
            logger.error("textract 套件未安裝，無法解析PDF。請執行 `pip install textract`")
            raise ImportError("處理PDF需要 textract 套件。 সন")
        try:
            byte_content = textract.process(file_path)
            content = byte_content.decode('utf-8')
            metadata = {'extractor': 'textract'}
            return content, metadata
        except Exception as e:
            logger.error(f"使用 textract 解析PDF失敗 {file_path}: {e}")
            raise e

    def _parse_docx(self, file_path: str) -> tuple[str, Dict]:
        """解析Word文檔 (需要安裝 python-docx)"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs])
            metadata = {'paragraphs': len(doc.paragraphs)}
            return content, metadata
        except ImportError:
            logger.warning("python-docx 未安裝，無法解析DOCX文件")
            raise ValueError("需要安裝 python-docx 來處理Word文檔")

# ... (The rest of the file remains the same)
class QuarterValidator:
    """季度格式驗證器"""
    
    @staticmethod
    def validate_quarter(quarter_str: str) -> bool:
        """驗證季度格式"""
        pattern = r'^\d{4}Q[1-4]$'
        return bool(re.match(pattern, quarter_str))


class BatchProcessor:
    """批量處理器"""
    
    def __init__(self, config: BatchProcessingConfig = None):
        self.config = config or BatchProcessingConfig()
        self.file_parser = FileParser()
        self.quarter_validator = QuarterValidator()
        
        # 處理統計
        self.processing_stats = {
            'total_files': 0,
            'successful_files': 0,
            'failed_files': 0,
            'total_chunks': 0,
            'total_entities': 0,
            'total_relations': 0,
            'processing_time': 0.0,
            'errors': []
        }
        
        # 進度追蹤
        self.processed_files = []
        self.failed_files = []
    
    def discover_files(self, directory: str, recursive: bool = True) -> List[str]:
        """發現目錄中的所有支持文件"""
        supported_exts = list(self.file_parser.supported_formats.keys())
        files = []
        
        path = Path(directory)
        if not path.exists():
            raise ValueError(f"目錄不存在: {directory}")
        
        # 遞歸或非遞歸搜索
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in supported_exts:
                files.append(str(file_path))
        
        logger.info(f"發現 {len(files)} 個支持的文件")
        return files
    
    def prepare_documents(self, file_paths: List[str], 
                         required_quarter: str) -> List[DocumentInfo]:
        """準備文檔信息（需要手動指定季度）"""
        if not required_quarter:
            raise ValueError("必須指定季度信息，不支持自動檢測")
        
        if not self.quarter_validator.validate_quarter(required_quarter):
            raise ValueError(f"無效的季度格式: {required_quarter}，應為 YYYYQN 格式（如 2024Q1）")
        
        documents = []
        
        for file_path in file_paths:
            try:
                # 解析文件
                parsed = self.file_parser.parse_file(file_path)
                
                # 生成文檔ID
                doc_id = Path(file_path).stem
                
                doc_info = DocumentInfo(
                    file_path=file_path,
                    doc_id=doc_id,
                    quarter=required_quarter,  # 使用手動指定的季度
                    content=parsed['content'],
                    file_size=parsed['file_size'],
                    file_format=parsed['file_format'],
                    metadata=parsed['metadata']
                )
                
                documents.append(doc_info)
                logger.info(f"準備文檔: {file_path} -> {required_quarter}")
                
            except Exception as e:
                error_msg = f"準備文檔失敗 {file_path}: {e}"
                logger.error(error_msg)
                self.processing_stats['errors'].append(error_msg)
                
                if not self.config.skip_on_error:
                    raise
        
        logger.info(f"成功準備 {len(documents)} 個文檔（季度: {required_quarter}）")
        return documents
    
    def prepare_documents_with_mapping(self, file_quarter_mapping: Dict[str, str]) -> List[DocumentInfo]:
        """使用手動指定的文件-季度映射準備文檔"""
        documents = []
        
        logger.info(f"使用手動映射準備 {len(file_quarter_mapping)} 個文檔")
        
        for file_path, quarter in file_quarter_mapping.items():
            try:
                # 驗證季度格式
                if not self.quarter_validator.validate_quarter(quarter):
                    logger.warning(f"無效的季度格式 '{quarter}' 用於文件 {file_path}")
                    continue
                
                # 解析文件
                parsed = self.file_parser.parse_file(file_path)
                
                # 生成文檔ID
                doc_id = Path(file_path).stem
                
                doc_info = DocumentInfo(
                    file_path=file_path,
                    doc_id=doc_id,
                    quarter=quarter,  # 使用手動指定的季度
                    content=parsed['content'],
                    file_size=parsed['file_size'],
                    file_format=parsed['file_format'],
                    metadata=parsed['metadata']
                )
                
                documents.append(doc_info)
                logger.info(f"準備文檔: {file_path} -> {quarter}")
                
            except Exception as e:
                error_msg = f"準備文檔失敗 {file_path}: {e}"
                logger.error(error_msg)
                self.processing_stats['errors'].append(error_msg)
                
                if not self.config.skip_on_error:
                    raise
        
        logger.info(f"成功準備 {len(documents)} 個文檔（手動映射）")
        return documents
    
    def process_batch(self, documents: List[DocumentInfo], 
                     graph_rag_system) -> Dict[str, Any]:
        """批量處理文檔"""
        start_time = time.time()
        self.processing_stats['total_files'] = len(documents)
        
        logger.info(f"開始批量處理 {len(documents)} 個文檔")
        
        # 選擇處理方式
        if self.config.max_workers <= 1:
            # 單線程處理
            results = self._process_sequential(documents, graph_rag_system)
        else:
            # 並行處理
            results = self._process_parallel(documents, graph_rag_system)
        
        # 完成圖建構
        graph_rag_system.finalize_graph()
        
        # 更新統計信息
        self.processing_stats['processing_time'] = time.time() - start_time
        
        logger.info(f"批量處理完成，耗時 {self.processing_stats['processing_time']:.2f}s")
        
        return {
            'results': results,
            'stats': self.processing_stats,
            'successful_files': self.processed_files,
            'failed_files': self.failed_files
        }
    
    def _process_sequential(self, documents: List[DocumentInfo], 
                          graph_rag_system) -> List[DocumentInfo]:
        """順序處理文檔"""
        results = []
        
        for i, doc in enumerate(documents, 1):
            logger.info(f"處理文檔 {i}/{len(documents)}: {doc.doc_id}")
            
            try:
                result = self._process_single_document(doc, graph_rag_system)
                results.append(result)
                self.processed_files.append(result)
                self.processing_stats['successful_files'] += 1
                
            except Exception as e:
                error_msg = f"處理文檔失敗 {doc.file_path}: {e}"
                logger.error(error_msg)
                
                doc.error = str(e)
                results.append(doc)
                self.failed_files.append(doc)
                self.processing_stats['failed_files'] += 1
                self.processing_stats['errors'].append(error_msg)
                
                if not self.config.skip_on_error:
                    raise
        
        return results
    
    def _process_parallel(self, documents: List[DocumentInfo], 
                         graph_rag_system) -> List[DocumentInfo]:
        """並行處理文檔"""
        results = []
        
        # 選擇執行器
        executor_class = (ProcessPoolExecutor if self.config.use_multiprocessing 
                         else ThreadPoolExecutor)
        
        with executor_class(max_workers=self.config.max_workers) as executor:
            # 提交任務
            future_to_doc = {
                executor.submit(self._process_single_document, doc, graph_rag_system): doc
                for doc in documents
            }
            
            # 收集結果
            completed = 0
            for future in as_completed(future_to_doc):
                doc = future_to_doc[future]
                completed += 1
                
                try:
                    result = future.result()
                    results.append(result)
                    self.processed_files.append(result)
                    self.processing_stats['successful_files'] += 1
                    
                    logger.info(f"完成 {completed}/{len(documents)}: {doc.doc_id}")
                    
                except Exception as e:
                    error_msg = f"處理文檔失敗 {doc.file_path}: {e}"
                    logger.error(error_msg)
                    
                    doc.error = str(e)
                    results.append(doc)
                    self.failed_files.append(doc)
                    self.processing_stats['failed_files'] += 1
                    self.processing_stats['errors'].append(error_msg)
                    
                    if not self.config.skip_on_error:
                        # 取消所有未完成的任務
                        for f in future_to_doc:
                            f.cancel()
                        raise
        
        return results
    
    def _process_single_document(self, doc: DocumentInfo, 
                               graph_rag_system) -> DocumentInfo:
        """處理單個文檔"""
        start_time = time.time()
        
        try:
            # 調用圖RAG系統處理文檔
            initial_stats = graph_rag_system.processing_stats.copy()
            graph_rag_system.process_document(doc.content, doc.quarter, doc.doc_id)
            
            # 計算增量統計
            new_stats = graph_rag_system.processing_stats
            doc.chunk_count = new_stats['total_chunks'] - initial_stats['total_chunks']
            doc.entity_count = new_stats['entities_extracted'] - initial_stats['entities_extracted']
            doc.relation_count = new_stats['relations_extracted'] - initial_stats['relations_extracted']
            
            # 更新批量統計
            self.processing_stats['total_chunks'] += doc.chunk_count
            self.processing_stats['total_entities'] += doc.entity_count
            self.processing_stats['total_relations'] += doc.relation_count
            
            doc.processing_time = time.time() - start_time
            
            logger.info(f"文檔處理成功 {doc.doc_id}: "
                       f"{doc.chunk_count} chunks, "
                       f"{doc.entity_count} entities, "
                       f"{doc.relation_count} relations")
            
            return doc
            
        except Exception as e:
            doc.processing_time = time.time() - start_time
            raise e
    
    def save_progress(self, filepath: str = None):
        """保存處理進度"""
        if not self.config.save_progress:
            return
        
        filepath = filepath or self.config.progress_file
        
        progress_data = {
            'timestamp': datetime.now().isoformat(),
            'stats': self.processing_stats,
            'processed_files': [
                {
                    'file_path': doc.file_path,
                    'doc_id': doc.doc_id,
                    'quarter': doc.quarter,
                    'processing_time': doc.processing_time,
                    'chunk_count': doc.chunk_count,
                    'entity_count': doc.entity_count,
                    'relation_count': doc.relation_count
                }
                for doc in self.processed_files
            ],
            'failed_files': [
                {
                    'file_path': doc.file_path,
                    'doc_id': doc.doc_id,
                    'error': doc.error
                }
                for doc in self.failed_files
            ]
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"進度已保存到: {filepath}")
    
    def get_processing_summary(self) -> Dict[str, Any]:
        """獲取處理摘要"""
        return {
            'total_files': self.processing_stats['total_files'],
            'successful_files': self.processing_stats['successful_files'],
            'failed_files': self.processing_stats['failed_files'],
            'success_rate': (self.processing_stats['successful_files'] / 
                           max(self.processing_stats['total_files'], 1)) * 100,
            'total_processing_time': self.processing_stats['processing_time'],
            'avg_time_per_file': (self.processing_stats['processing_time'] / 
                                max(self.processing_stats['successful_files'], 1)),
            'total_chunks': self.processing_stats['total_chunks'],
            'total_entities': self.processing_stats['total_entities'],
            'total_relations': self.processing_stats['total_relations'],
            'errors_count': len(self.processing_stats['errors'])
        }
